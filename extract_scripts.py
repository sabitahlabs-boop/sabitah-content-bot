"""
Extract content from Google Docs that may be imported .docx files.

Tries 3 strategies in order:
  1. Google Docs API documents().get() — read body.content
  2. Google Drive API files().export(mimeType='text/plain') — download as plain text
  3. Google Drive API files().get_media() — download original .docx + parse with python-docx

Usage:
    py extract_scripts.py
"""
import io
import json
import os
import sys
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from telegram_bot import get_google_credentials


# Documents to extract: doc_id → label
DOCS_TO_EXTRACT = {
    "1YRgFI52HKXJkM1d0zJNazF_wLup5kFzzkj5a8suqVsI": '[SB-008] Sabitah - "Kenapa caption kamu cuma dibaca tapi nggak di-klik?"',
    "15Cumro1zOrcmJ2R2hnlFKhr8Baqdd67_3uCOJP1gR94": "[CT-027] County - Pembukuan UMKM",
    "1i6bnw0A71UWhwLhFB91UwgWg3pCtJwXZ-QAFp2Vb-8A": "[LG-024] LEGUS - Hak hukum pengusaha",
    "13jUvt8Zm9-tE1WsuNRmYXjjxZ3xcE07bxlVEGXMoApE": "[DF-026] Defarchy - Sepeda listrik commuting",
    "13AQ2MFY4jElUNNfDcr46ZcWmE6V6mwD8v3cOCVIv8IU": "[HB-010] Happy Baby - #HappyBabySteps",
}

OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "extracted_scripts.json")


def strategy_1_docs_api(docs_service, doc_id):
    """Strategy 1: Google Docs API documents().get() — only works for native Google Docs."""
    try:
        doc = docs_service.documents().get(documentId=doc_id).execute()
        text = ""
        for element in doc.get("body", {}).get("content", []):
            if "paragraph" in element:
                for elem in element["paragraph"].get("elements", []):
                    if "textRun" in elem:
                        text += elem["textRun"]["content"]
            elif "table" in element:
                for row in element["table"].get("tableRows", []):
                    for cell in row.get("tableCells", []):
                        for c in cell.get("content", []):
                            if "paragraph" in c:
                                for elem in c["paragraph"].get("elements", []):
                                    if "textRun" in elem:
                                        text += elem["textRun"]["content"]
        return text.strip(), doc.get("title", "")
    except Exception as e:
        return "", f"ERROR: {e}"


def strategy_2_drive_export(drive_service, doc_id):
    """Strategy 2: Drive API export to text/plain — works for native Google Docs."""
    try:
        request = drive_service.files().export_media(
            fileId=doc_id, mimeType="text/plain"
        )
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        return buf.getvalue().decode("utf-8", errors="replace").strip(), ""
    except Exception as e:
        return "", f"ERROR: {e}"


def strategy_3_download_docx(drive_service, doc_id):
    """Strategy 3: Download original file via get_media() + parse .docx with python-docx."""
    try:
        # Get file metadata to know the mime type
        meta = drive_service.files().get(
            fileId=doc_id, fields="name,mimeType"
        ).execute()
        mime = meta.get("mimeType", "")
        name = meta.get("name", "")

        # Download the file
        request = drive_service.files().get_media(fileId=doc_id)
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        buf.seek(0)

        # Parse based on mime type
        if "wordprocessingml" in mime or name.lower().endswith(".docx"):
            try:
                from docx import Document
            except ImportError:
                return "", "ERROR: python-docx not installed. Run: pip install python-docx"

            doc = Document(buf)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n".join(paragraphs).strip(), f"mime={mime}, name={name}"

        # If it's a Google Doc that wasn't exported successfully, fallback to raw
        return buf.getvalue().decode("utf-8", errors="replace").strip(), f"mime={mime}, name={name}"
    except Exception as e:
        return "", f"ERROR: {e}"


def extract_doc(docs_service, drive_service, doc_id, label):
    """Try all 3 strategies in order. Return (content, strategy_used, info)."""
    print(f"\n{'='*70}")
    print(f"DOC: {label}")
    print(f"ID:  {doc_id}")
    print(f"{'='*70}")

    # Strategy 1
    print("\n[Strategy 1] Google Docs API documents().get()...")
    content, info = strategy_1_docs_api(docs_service, doc_id)
    if content:
        print(f"  SUCCESS — {len(content)} chars (title: {info})")
        return content, "docs_api", info
    print(f"  FAILED — content empty. {info}")

    # Strategy 2
    print("\n[Strategy 2] Drive API export to text/plain...")
    content, info = strategy_2_drive_export(drive_service, doc_id)
    if content:
        print(f"  SUCCESS — {len(content)} chars")
        return content, "drive_export", info
    print(f"  FAILED — {info}")

    # Strategy 3
    print("\n[Strategy 3] Download original file + parse with python-docx...")
    content, info = strategy_3_download_docx(drive_service, doc_id)
    if content:
        print(f"  SUCCESS — {len(content)} chars ({info})")
        return content, "download_docx", info
    print(f"  FAILED — {info}")

    return "", "all_failed", info


def main():
    print("=" * 70)
    print("EXTRACT SCRIPTS — Multi-Strategy Google Docs Reader")
    print("=" * 70)

    creds = get_google_credentials()
    docs_service = build("docs", "v1", credentials=creds)
    drive_service = build("drive", "v3", credentials=creds)

    results = {}
    summary = {"docs_api": 0, "drive_export": 0, "download_docx": 0, "all_failed": 0}

    for doc_id, label in DOCS_TO_EXTRACT.items():
        content, strategy, info = extract_doc(docs_service, drive_service, doc_id, label)
        summary[strategy] += 1

        results[doc_id] = {
            "label": label,
            "strategy_used": strategy,
            "info": info,
            "content_length": len(content),
            "content": content,
        }

        # Print full content
        print(f"\n--- CONTENT ({len(content)} chars) ---")
        if content:
            print(content)
        else:
            print("(EMPTY — all strategies failed)")
        print(f"--- END ---")

    # Save to JSON
    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)

    print(f"\n\n{'=' * 70}")
    print("SUMMARY")
    print(f"{'=' * 70}")
    print(f"Total docs: {len(DOCS_TO_EXTRACT)}")
    print(f"  Strategy 1 (Docs API):    {summary['docs_api']}")
    print(f"  Strategy 2 (Drive Export): {summary['drive_export']}")
    print(f"  Strategy 3 (Download .docx): {summary['download_docx']}")
    print(f"  All failed:               {summary['all_failed']}")
    print(f"\nResults saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
